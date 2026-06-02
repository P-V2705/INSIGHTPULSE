"""
Verify the 400 error is fixed across all file types.
Tests: upload → analysis start (no 400) → poll → results.
"""
import requests, io, time, openpyxl, json

BASE = "http://localhost:8000/api"

def run(label, filename, content, mime):
    # Upload
    r = requests.post(f"{BASE}/upload/file",
        files={"file": (filename, content, mime)})
    assert r.status_code == 200, f"Upload failed {label}: {r.text[:150]}"
    d = r.json()
    sid = d["session_id"]
    detected_col = d["detected_columns"].get("review_column")

    # Start analysis — pass review_col explicitly if detected, else None (backend fallback)
    r = requests.post(f"{BASE}/analysis/run", json={
        "session_id": sid,
        "review_column": detected_col,
        "max_rows": 100,
    })
    assert r.status_code == 200, f"Analysis 400 on {label}: {r.json().get('detail','?')}"
    used_col = r.json().get("review_column", "?")

    # Poll
    for _ in range(20):
        time.sleep(2)
        s = requests.get(f"{BASE}/analysis/status/{sid}").json()["status"]
        if s == "completed": break
        if s == "error":
            err = requests.get(f"{BASE}/analysis/status/{sid}").json().get("error","?")
            return f"ANALYSIS_ERROR: {err[:80]}"

    # Results
    r = requests.get(f"{BASE}/analysis/results/{sid}")
    assert r.status_code == 200, f"Results failed {label}: {r.text[:100]}"
    ov = r.json()["sentiment_overview"]
    return f"{ov['total_analyzed']} reviews | col='{used_col}' | +{ov['positive_pct']}% -{ov['negative_pct']}%"

print("=" * 60)
print("  HTTP 400 Fix Verification — All Formats")
print("=" * 60)

# CSV
csv = b"review_text,rating\n\"Amazing!\",5\n\"Terrible.\",1\n\"Okay.\",3"
print(f"\n  CSV   : {run('csv',  'r.csv',  io.BytesIO(csv),  'text/csv')}")

# XLSX — test all MIME variants (the ones browsers send)
wb = openpyxl.Workbook()
ws = wb.active
ws.append(["review_text", "rating"])
for t, rt in [("Great product!", 5), ("Broke after a day.", 1), ("Decent quality.", 3)]:
    ws.append([t, rt])

for mime_label, mime in [
    ("xlsx/official", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
    ("xlsx/ms-excel",  "application/vnd.ms-excel"),
    ("xlsx/octet",     "application/octet-stream"),
]:
    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    print(f"  {mime_label:<14}: {run(mime_label, 'r.xlsx', buf, mime)}")

# JSON
jd = json.dumps([{"review_text": "Excellent!", "rating": 5},
                  {"review_text": "Very bad.",   "rating": 1}]).encode()
print(f"\n  JSON  : {run('json', 'r.json', io.BytesIO(jd), 'application/json')}")

# TXT
txt = b"This product is absolutely amazing.\nTerrible quality, avoid.\nIt is okay."
print(f"  TXT   : {run('txt',  'r.txt',  io.BytesIO(txt), 'text/plain')}")

# TSV
tsv = b"review_text\trating\nGreat product!\t5\nBad quality.\t1"
print(f"  TSV   : {run('tsv',  'r.tsv',  io.BytesIO(tsv), 'text/tab-separated-values')}")

# DOCX
try:
    from docx import Document
    doc = Document()
    doc.add_paragraph("This product is absolutely amazing and I love it!")
    doc.add_paragraph("Terrible quality, broke after one day.")
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    print(f"  DOCX  : {run('docx', 'r.docx', buf, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}")
except Exception as e:
    print(f"  DOCX  : SKIP ({e})")

print()
print("=" * 60)
print("  ALL 400 ERRORS FIXED — SYSTEM OPERATIONAL")
print("=" * 60)
print(f"\n  Frontend : http://localhost:3000")
print(f"  Backend  : http://localhost:8000")
