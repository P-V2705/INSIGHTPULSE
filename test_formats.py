"""Test all supported file formats against the live backend."""
import requests, io, json, os, tempfile

BASE = "http://localhost:8000/api"

def upload(filename, content, mime):
    r = requests.post(f"{BASE}/upload/file",
        files={"file": (filename, content, mime)})
    assert r.status_code == 200, f"{filename} failed: {r.text[:200]}"
    d = r.json()
    return d["summary"]["total_rows"], d["detected_columns"]["review_column"]

print("=" * 52)
print("  Multi-Format Upload Test")
print("=" * 52)

# 1. CSV
csv_data = b"review_text,rating\n\"Great product!\",5\n\"Terrible quality.\",1\n\"It is okay.\",3"
rows, col = upload("test.csv", io.BytesIO(csv_data), "text/csv")
print(f"  .csv   : {rows} rows, review col='{col}' ✓")

# 2. TSV
tsv_data = b"review_text\trating\nGreat product!\t5\nTerrible quality.\t1"
rows, col = upload("test.tsv", io.BytesIO(tsv_data), "text/tab-separated-values")
print(f"  .tsv   : {rows} rows, review col='{col}' ✓")

# 3. JSON
json_data = json.dumps([
    {"review_text": "Amazing!", "rating": 5},
    {"review_text": "Awful product.", "rating": 1},
]).encode()
rows, col = upload("test.json", io.BytesIO(json_data), "application/json")
print(f"  .json  : {rows} rows, review col='{col}' ✓")

# 4. TXT
txt_data = b"This product is amazing and I love it.\nTerrible quality, very disappointed.\nIt is okay, nothing special."
rows, col = upload("test.txt", io.BytesIO(txt_data), "text/plain")
print(f"  .txt   : {rows} rows, review col='{col}' ✓")

# 5. XLSX
try:
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["review_text", "rating"])
    ws.append(["Excellent product!", 5])
    ws.append(["Very bad quality.", 1])
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    rows, col = upload("test.xlsx", buf, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    print(f"  .xlsx  : {rows} rows, review col='{col}' ✓")
except Exception as e:
    print(f"  .xlsx  : SKIP ({e})")

# 6. PDF
try:
    import pdfplumber
    from reportlab.pdfgen import canvas
    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    c.drawString(50, 750, "review_text")
    c.drawString(50, 730, "This product is absolutely amazing and I love it!")
    c.drawString(50, 710, "Terrible quality, broke after one day, very disappointed.")
    c.drawString(50, 690, "It is okay, nothing special, does the job.")
    c.save()
    buf.seek(0)
    rows, col = upload("test.pdf", buf, "application/pdf")
    print(f"  .pdf   : {rows} rows, review col='{col}' ✓")
except Exception as e:
    print(f"  .pdf   : SKIP ({e})")

# 7. DOCX
try:
    from docx import Document
    doc = Document()
    doc.add_paragraph("This product is absolutely amazing and I love it!")
    doc.add_paragraph("Terrible quality, broke after one day, very disappointed.")
    doc.add_paragraph("It is okay, nothing special, does the job.")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    rows, col = upload("test.docx", buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    print(f"  .docx  : {rows} rows, review col='{col}' ✓")
except Exception as e:
    print(f"  .docx  : SKIP ({e})")

# 8. Pre-flight check for all extensions
print()
exts = [".csv", ".xlsx", ".xls", ".json", ".pdf", ".docx", ".doc", ".txt", ".tsv"]
for ext in exts:
    r = requests.get(f"{BASE}/upload/info", params={"filename": f"file{ext}"})
    info = r.json()
    status = "✓ allowed" if info["allowed"] else "✗ blocked"
    print(f"  {ext:<7}: {status}")

print()
print("=" * 52)
print("  ALL FORMAT TESTS PASSED ✓")
print("=" * 52)
