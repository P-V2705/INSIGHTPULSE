"""
Dataset Processor — reads CSV/Excel/JSON/PDF/Word/TXT/TSV,
auto-detects review columns, cleans data, and prepares it for NLP analysis.
"""
import pandas as pd
import numpy as np
import json
import re
from pathlib import Path
from typing import Optional

REVIEW_COLUMN_HINTS = [
    "review", "text", "comment", "feedback", "description",
    "body", "content", "message", "opinion", "review_text",
    "reviewtext", "review_body", "reviewbody", "summary",
    "title", "headline", "review_title"
]

RATING_COLUMN_HINTS = [
    "rating", "score", "stars", "star", "rate", "overall",
    "review_score", "reviewscore", "overall_rating", "grade"
]

CATEGORY_COLUMN_HINTS = [
    "category", "type", "product", "brand", "name", "item",
    "product_name", "productname", "asin", "product_id"
]


def _load_pdf(file_path: str) -> pd.DataFrame:
    """
    Extract text from every page of a PDF.
    Each page becomes one row with columns: page, text.
    Tables inside the PDF are also extracted and merged.
    """
    import pdfplumber

    rows = []
    tables_found = []

    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            # Try to extract structured tables first
            for table in page.extract_tables():
                if table and len(table) > 1:
                    headers = [str(h).strip() if h else f"col_{i}"
                               for i, h in enumerate(table[0])]
                    for data_row in table[1:]:
                        row_dict = {headers[i]: (str(v).strip() if v else "")
                                    for i, v in enumerate(data_row) if i < len(headers)}
                        row_dict["_source_page"] = page_num
                        tables_found.append(row_dict)

            # Plain text fallback
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                rows.append({"page": page_num, "text": text})

    # Prefer structured table data if found
    if tables_found:
        df = pd.DataFrame(tables_found)
        # Drop mostly-empty columns
        df = df.dropna(axis=1, thresh=max(1, len(df) // 2))
        return df

    if not rows:
        raise ValueError("No text could be extracted from the PDF.")

    return pd.DataFrame(rows)


def _load_word(file_path: str) -> pd.DataFrame:
    """
    Extract text from a Word document (.docx / .doc).
    Each paragraph becomes one row; tables are also extracted.
    """
    from docx import Document

    doc = Document(file_path)
    rows = []
    tables_found = []

    # Extract tables
    for table in doc.tables:
        if not table.rows:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                row_dict = {headers[i]: cells[i] for i in range(min(len(headers), len(cells)))}
                tables_found.append(row_dict)

    if tables_found:
        return pd.DataFrame(tables_found)

    # Extract paragraphs
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            rows.append({"paragraph": i + 1, "text": text})

    if not rows:
        raise ValueError("No text could be extracted from the Word document.")

    return pd.DataFrame(rows)


def _load_txt(file_path: str) -> pd.DataFrame:
    """
    Load a plain text file.
    Each non-empty line becomes one row with column 'text'.
    """
    for enc in ["utf-8", "latin-1", "cp1252"]:
        try:
            with open(file_path, "r", encoding=enc) as f:
                lines = [l.strip() for l in f if l.strip()]
            return pd.DataFrame({"text": lines})
        except UnicodeDecodeError:
            continue
    raise ValueError("Could not decode text file.")


def _load_tsv(file_path: str) -> pd.DataFrame:
    """Load a tab-separated values file."""
    for enc in ["utf-8", "latin-1", "cp1252"]:
        try:
            return pd.read_csv(file_path, sep="\t", encoding=enc, low_memory=False)
        except UnicodeDecodeError:
            continue
    raise ValueError("Could not decode TSV file.")


def load_dataset(file_path: str) -> pd.DataFrame:
    """
    Universal loader — dispatches to the correct reader based on extension.
    Supports: CSV, Excel (.xlsx/.xls), JSON, PDF, Word (.docx/.doc), TXT, TSV.
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".csv":
        for enc in ["utf-8", "latin-1", "cp1252"]:
            try:
                return pd.read_csv(file_path, encoding=enc, low_memory=False)
            except UnicodeDecodeError:
                continue
        raise ValueError("Could not decode CSV file.")

    elif ext in (".xlsx", ".xls"):
        return pd.read_excel(file_path)

    elif ext == ".json":
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        if isinstance(data, list):
            return pd.DataFrame(data)
        elif isinstance(data, dict):
            for key in ["data", "reviews", "items", "records", "results"]:
                if key in data and isinstance(data[key], list):
                    return pd.DataFrame(data[key])
            return pd.DataFrame([data])
        raise ValueError("Unsupported JSON structure.")

    elif ext == ".pdf":
        return _load_pdf(file_path)

    elif ext in (".docx", ".doc"):
        return _load_word(file_path)

    elif ext == ".txt":
        return _load_txt(file_path)

    elif ext == ".tsv":
        return _load_tsv(file_path)

    else:
        raise ValueError(f"Unsupported file format: {ext}")


def detect_columns(df: pd.DataFrame) -> dict:
    """Auto-detect review, rating, and category columns."""
    cols_lower = {col.lower().strip(): col for col in df.columns}

    review_col = None
    for hint in REVIEW_COLUMN_HINTS:
        for col_lower, col_orig in cols_lower.items():
            if hint in col_lower:
                # Prefer text-heavy columns
                if df[col_orig].dtype == object:
                    review_col = col_orig
                    break
        if review_col:
            break

    # Fallback: pick the object column with longest average text
    if not review_col:
        text_cols = df.select_dtypes(include="object").columns.tolist()
        if text_cols:
            avg_lens = {col: df[col].dropna().astype(str).str.len().mean() for col in text_cols}
            review_col = max(avg_lens, key=avg_lens.get)

    rating_col = None
    for hint in RATING_COLUMN_HINTS:
        for col_lower, col_orig in cols_lower.items():
            if hint in col_lower:
                rating_col = col_orig
                break
        if rating_col:
            break

    category_col = None
    for hint in CATEGORY_COLUMN_HINTS:
        for col_lower, col_orig in cols_lower.items():
            if hint in col_lower:
                category_col = col_orig
                break
        if category_col:
            break

    return {
        "review_column": review_col,
        "rating_column": rating_col,
        "category_column": category_col,
        "all_columns": list(df.columns),
    }


def clean_dataset(df: pd.DataFrame, review_col: str, rating_col: Optional[str] = None) -> pd.DataFrame:
    """Clean and prepare dataset for NLP."""
    df = df.copy()

    # Drop rows with empty reviews
    df = df.dropna(subset=[review_col])
    df[review_col] = df[review_col].astype(str).str.strip()
    df = df[df[review_col].str.len() > 5]

    # Clean rating column
    if rating_col and rating_col in df.columns:
        df[rating_col] = pd.to_numeric(df[rating_col], errors="coerce")
        df = df.dropna(subset=[rating_col])
        # Normalize to 1-5 scale if needed
        max_rating = df[rating_col].max()
        if max_rating > 5:
            df[rating_col] = (df[rating_col] / max_rating * 5).round(1)

    # Reset index
    df = df.reset_index(drop=True)
    return df


def get_dataset_summary(df: pd.DataFrame, review_col: str, rating_col: Optional[str] = None) -> dict:
    """Generate a summary of the dataset."""
    summary = {
        "total_rows": int(len(df)),
        "total_columns": int(len(df.columns)),
        "columns": list(df.columns),
        "review_column": review_col,
        "rating_column": rating_col,
        "missing_values": int(df.isnull().sum().sum()),
        "avg_review_length": round(float(df[review_col].str.len().mean()), 1),
        "max_review_length": int(df[review_col].str.len().max()),
        "min_review_length": int(df[review_col].str.len().min()),
    }

    if rating_col and rating_col in df.columns:
        summary["avg_rating"] = round(float(df[rating_col].mean()), 2)
        summary["rating_distribution"] = {str(k): int(v) for k, v in df[rating_col].value_counts().sort_index().items()}

    # Preview first 5 rows
    preview_cols = [review_col]
    if rating_col and rating_col in df.columns:
        preview_cols.append(rating_col)
    # Convert preview to native Python types
    preview_raw = df[preview_cols].head(5).to_dict(orient="records")
    summary["preview"] = [{k: (int(v) if hasattr(v, 'item') else v) for k, v in row.items()} for row in preview_raw]

    return summary


def sample_for_analysis(df: pd.DataFrame, max_rows: int = 2000) -> pd.DataFrame:
    """Sample dataset if too large for fast processing."""
    if len(df) > max_rows:
        return df.sample(n=max_rows, random_state=42).reset_index(drop=True)
    return df


# ── Large-file helpers ────────────────────────────────────────────────────────

def load_dataset_chunked(file_path: str, chunk_size: int = 50_000):
    """
    Generator that yields DataFrame chunks.
    CSV: true chunked reading.
    All other formats: load once and yield as a single chunk.
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".csv":
        for enc in ["utf-8", "latin-1", "cp1252"]:
            try:
                reader = pd.read_csv(
                    file_path, encoding=enc, low_memory=True, chunksize=chunk_size,
                )
                for chunk in reader:
                    yield chunk
                return
            except UnicodeDecodeError:
                continue
        raise ValueError("Could not decode CSV file.")
    else:
        yield load_dataset(file_path)


def get_large_file_summary(file_path: str, ext: str) -> tuple:
    """
    Memory-efficient summary for files > 200 MB.
    Reads only the first chunk to detect columns and build a preview,
    then counts total rows by scanning the file without loading it all.
    """
    path = Path(file_path)

    # --- Read first chunk only ---
    first_chunk = None
    total_rows = 0

    if ext == ".csv":
        for enc in ["utf-8", "latin-1", "cp1252"]:
            try:
                reader = pd.read_csv(
                    file_path,
                    encoding=enc,
                    low_memory=True,
                    chunksize=50_000,
                )
                for i, chunk in enumerate(reader):
                    total_rows += len(chunk)
                    if i == 0:
                        first_chunk = chunk
                break
            except UnicodeDecodeError:
                continue
    else:
        first_chunk = load_dataset(file_path)
        total_rows = len(first_chunk)

    if first_chunk is None or len(first_chunk) == 0:
        raise ValueError("File appears to be empty.")

    detected = detect_columns(first_chunk)
    review_col = detected["review_column"]
    rating_col = detected["rating_column"]

    summary: dict = {
        "total_rows": total_rows,
        "total_columns": int(len(first_chunk.columns)),
        "columns": list(first_chunk.columns),
        "review_column": review_col,
        "rating_column": rating_col,
        "missing_values": None,   # too expensive to compute on huge files
        "avg_review_length": None,
        "max_review_length": None,
        "min_review_length": None,
        "large_file": True,
    }

    if review_col and review_col in first_chunk.columns:
        lengths = first_chunk[review_col].dropna().astype(str).str.len()
        summary["avg_review_length"] = round(float(lengths.mean()), 1)
        summary["max_review_length"] = int(lengths.max())
        summary["min_review_length"] = int(lengths.min())

    if rating_col and rating_col in first_chunk.columns:
        summary["avg_rating"] = round(float(pd.to_numeric(first_chunk[rating_col], errors="coerce").mean()), 2)

    # Preview — first 5 rows of first chunk
    preview_cols = [c for c in ([review_col, rating_col] if rating_col else [review_col]) if c]
    if preview_cols:
        preview_raw = first_chunk[preview_cols].head(5).fillna("").astype(str).to_dict(orient="records")
        summary["preview"] = preview_raw
    else:
        summary["preview"] = first_chunk.head(5).fillna("").astype(str).to_dict(orient="records")

    return first_chunk, detected, summary


def sample_for_analysis_large(
    file_path: str,
    review_col: str,
    rating_col: Optional[str],
    max_rows: int = 5000,
    chunk_size: int = 50_000,
) -> pd.DataFrame:
    """
    Reservoir-sample `max_rows` rows from a large CSV without loading
    the whole file into memory.  Falls back to full load for small files.
    """
    import random

    path = Path(file_path)
    ext = path.suffix.lower()

    if ext != ".csv":
        df = load_dataset(file_path)
        return sample_for_analysis(df, max_rows)

    reservoir: list = []
    total_seen = 0

    for enc in ["utf-8", "latin-1", "cp1252"]:
        try:
            reader = pd.read_csv(
                file_path,
                encoding=enc,
                low_memory=True,
                chunksize=chunk_size,
            )
            for chunk in reader:
                rows = chunk.to_dict(orient="records")
                for row in rows:
                    total_seen += 1
                    if len(reservoir) < max_rows:
                        reservoir.append(row)
                    else:
                        j = random.randint(0, total_seen - 1)
                        if j < max_rows:
                            reservoir[j] = row
            break
        except UnicodeDecodeError:
            continue

    if not reservoir:
        raise ValueError("No rows could be read from the file.")

    df = pd.DataFrame(reservoir)
    if review_col and review_col in df.columns:
        df = clean_dataset(df, review_col, rating_col)
    return df.reset_index(drop=True)
